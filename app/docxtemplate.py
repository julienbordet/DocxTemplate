#!/usr/bin/env python3
import sys

import openpyxl
import docx
from docx.shared import Pt
import os
import logging
import shutil
import argparse
import re


class DocxTemplate(object):

    def __init__(self, model_name, xlsx_name, target_dir, prefix):
        self.model_name = model_name
        self.xlsx_name = xlsx_name
        self.prefix = prefix

        if not os.path.isdir(target_dir):
            logging.error(f"Directory does not exist : '{target_dir}'")
            raise NotADirectoryError
        self.target_dir = target_dir

        # Open docx template
        try:
            self.model_doc = docx.Document(self.model_name)
        except Exception as e:
            logging.error(f"Unable to open file : '{ self.model_name }' -> { e }")
            sys.exit(1)

        # Open Excel data file

        if not os.path.isfile(self.xlsx_name):
            logging.error(f"xlsx file unavailable : '{self.xlsx_name}'")
            raise FileNotFoundError

        try:
            self.data_xlsx = openpyxl.load_workbook(self.xlsx_name)
        except Exception as e:
            logging.error(f"Unable to open file : '{ self.xlsx_name }' -> { e }")
            sys.exit(1)

    def get_info(self):
        sheet = self.data_xlsx.active
        self.data = {}

        # Check for relevant columns
        j = 1
        columns = [""]
        while sheet.cell(row=1, column=j).value is not None:
            columns.append(sheet.cell(row=1, column=j).value)
            j += 1

        i = 2
        while sheet.cell(row=i, column=1).value is not None:
            data = {}

            key = sheet.cell(row=i, column=1).value
            if key in data.keys():
                logging.error("Some rows in the Excel file have the same keys. "
                              "Make sure first column values are unique")
                sys.exit(1)

            for col in range(1, j):
                v = sheet.cell(row=i, column=col).value
                if v is not None:
                    data[columns[col]] = str(v)
                else:
                    data[columns[col]] = None

            self.data[key] = data

            i += 1

    def generate_docx(self, key, item_data):
        target_name = f"{self.target_dir}/{self.prefix}-{key}.docx"
        shutil.copy(self.model_name, target_name)

        try:
            target_doc = docx.Document(target_name)
        except Exception as e:
            logging.error(f"Unable to open target file : '{target_name}' -> {e}")
            sys.exit(1)

        style = target_doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)

        for paragraph in target_doc.paragraphs:
            for variable in item_data:
                if f"***{variable}***" in paragraph.text:
                    logging.debug(f"{key} -> {variable} = {item_data[variable]}")

                    if item_data[variable] is not None:
                        paragraph.text = paragraph.text.replace(f"***{variable}***", item_data[variable])
                    else:
                        pd = paragraph._element
                        pd.getparent().remove(pd)

            paragraph.style = target_doc.styles['Normal']

        logging.info(f"Writing {target_name} for {key}")
        target_doc.save(target_name)

    def generate_docs(self):
        if not self.data:
            logging.error("Generating letters but not data from template")
            sys.exit(1)

        logging.info(f"Generating {len(self.data)} letter(s)")

        for item_key in self.data:
            item_data = self.data[item_key]

            self.generate_docx(item_key, item_data)


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('directory', help="target directory for generated files")
    parser.add_argument('xlsx_data', help="excel data file")
    parser.add_argument('template', help="docx template")
    parser.add_argument('-prefix', help="prefix for generated files")
    parser.add_argument("-v", help="verbose mode (-vv for more)", action="count")
    args = parser.parse_args()

    if args.v == 1:
        logging.basicConfig(level=logging.INFO)

    if args.v > 1:
        logging.basicConfig(level=logging.DEBUG)

    if not args.prefix:
        p = re.compile(r'.*/(.*)\..+')
        prefix = p.match(args.template).group(1)
    else:
        prefix = args.prefix

    dt = DocxTemplate(args.template, args.xlsx_data, args.directory, prefix)
    dt.get_info()
    dt.generate_docs()
