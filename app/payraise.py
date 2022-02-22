#!/usr/bin/env python3
import sys

import openpyxl
import docx
from docx.shared import Pt
import os
import logging
import shutil

import config

target_dir = f"target/{config.year}"
raise_xlsx_name = f"Augmentation-Primes-{int(config.year)-1}-{config.year}.xlsx"
docx_model_name = f"templates/Courriers-augmentation-{int(config.year)-1}-{config.year}.docx"


class PayRaise(object):
    def __init__(self, model_name, target_dir):
        self.model_name = model_name

        if not os.path.isdir(target_dir):
            logging.error(f"Directory does not exist : '{target_dir}'")
            raise NotADirectoryError
        self.target_dir = target_dir
        self.setup()

    def setup(self):
        try:
            self.model_doc = docx.Document(self.model_name)
        except Exception as e:
            logging.error(f"Unable to open file : '{ self.model_name }' -> { e }")
            sys.exit(1)

        xlsx_full_path_name = f"{target_dir}/{raise_xlsx_name}"

        if not os.path.isfile(xlsx_full_path_name):
            logging.error(f"xlsx file unavailable : '{xlsx_full_path_name}'")
            raise FileNotFoundError

        try:
            self.raise_xlsx = openpyxl.load_workbook(xlsx_full_path_name)
        except Exception as e:
            logging.error(f"Unable to open file : '{ xlsx_full_path_name }' -> { e }")
            sys.exit(1)

    def get_raise_info(self):
        sheet = self.raise_xlsx.active
        self.data = {}

        i = 2
        while sheet.cell(row=i, column=1).value is not None:
            name = sheet.cell(row = i, column = 1).value
            first_name = sheet.cell(row = i, column = 2).value
            bonus = sheet.cell(row = i, column = 3).value
            payraise = sheet.cell(row = i, column = 4).value

            if bonus or payraise:
                if bonus is None:
                    bonus = 0

                if payraise is None:
                    payraise = 0

                self.data[f"{name}-{first_name}"] = {
                    'name': name,
                    'first_name': first_name,
                    'bonus' : str(round(bonus)),
                    'payraise' : str(payraise)
                }

            i += 1

    def generate_letter(self, first_name, name, bonus, payraise):
        target_name = f"{self.target_dir}/{config.year}-COURRIER-AUGMENTATION-{first_name}-{name}.docx"
        shutil.copy(self.model_name, target_name)

        try:
            target_doc = docx.Document(target_name)
        except Exception as e:
            logging.error(f"Unable to open target file : '{target_name}' -> {e}")
            sys.exit(1)

        style = target_doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)

        for paragraph in target_doc.paragraphs:
            if "***Prenom***" in paragraph.text:
                paragraph.text = paragraph.text.replace("***Prenom***", first_name)
            if "***Nom***" in paragraph.text:
                paragraph.text = paragraph.text.replace("***Nom***", name)
            if "***Prime***" in paragraph.text:
                if bonus != "0":
                    paragraph.text = paragraph.text.replace("***Prime***", bonus)
                else:
                    pd = paragraph._element
                    pd.getparent().remove(pd)

            if "***Augmentation***" in paragraph.text:
                if payraise != "0":
                    paragraph.text = paragraph.text.replace("***Augmentation***", payraise)
                else:
                    pd = paragraph._element
                    pd.getparent().remove(pd)

            paragraph.style = target_doc.styles['Normal']

        target_doc.save(target_name)

    def generate_letters(self):
        for employee in self.data:
            employee_data = self.data[employee]

            if employee_data['bonus'] or employee_data['payraise']:
                self.generate_letter(employee_data['first_name'], employee_data['name'],
                                       employee_data['bonus'], employee_data['payraise'])


if __name__ == '__main__':
    p = PayRaise(docx_model_name, target_dir)
    p.get_raise_info()
    p.generate_letters()
